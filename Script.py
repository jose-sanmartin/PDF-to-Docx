import PyMuPDF
import fitz
from docx import Document
from docx.shared import Pt
import unicodedata

def clean_text(text):
    """
    Limpia el texto eliminando caracteres que no son compatibles con XML.
    Esto incluye control characters y NULL bytes.
    """
    text = ''.join(char for char in text if char.isprintable())
    text = text.replace('\x00', '')  # Eliminar NULL bytes
    text = unicodedata.normalize('NFKC', text)  # Normalizar el texto
    return text

def convert_pdf_to_docx_with_styles(pdf_path, docx_path):
    pdf = fitz.open(pdf_path)
    doc = Document()

    for page_num in range(len(pdf)):
        page = pdf.load_page(page_num)
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if b["type"] == 0:  # Bloque de texto
                for line in b["lines"]:
                    for span in line["spans"]:
                        text = clean_text(span["text"])  # Limpieza de texto
                        size = span["size"]
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(text)
                        run.font.size = Pt(size)  # Aplicar el tamaño de fuente

    doc.save(docx_path)

# Rutas al archivo PDF y al archivo DOCX de salida
pdf_path = 'direction_file.pdf'
docx_path = 'direction_file.docx'
convert_pdf_to_docx_with_styles(pdf_path, docx_path)
