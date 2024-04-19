import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
import unicodedata
import os

def clean_text(text):
    text = ''.join(char for char in text if char.isprintable())
    text = text.replace('\x00', '')  # Eliminar NULL bytes
    text = unicodedata.normalize('NFKC', text)  # Normalizar el texto
    return text

def convert_pdf_to_docx_with_styles(pdf_path, docx_path):
    pdf = fitz.open(pdf_path)
    doc = Document()

    for page_num in range(len(pdf)):
        page = pdf.load_page(page_num)
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if b["type"] == 0:  # Bloque de texto
                for line in b["lines"]:
                    for span in line["spans"]:
                        text = clean_text(span["text"])  # Limpieza de texto
                        size = span["size"]
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(text)
                        run.font.size = Pt(size)  # Aplicar el tamaño de fuente

    doc.save(docx_path)

def convert_all_pdfs_in_folder(folder_path):
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(folder_path, filename)
            docx_path = os.path.join(folder_path, os.path.splitext(filename)[0] + '.docx')
            convert_pdf_to_docx_with_styles(pdf_path, docx_path)
            print(f"Convertido: {pdf_path} a {docx_path}")

# Especifique la ruta de la carpeta que contiene los archivos PDF
folder_path = 'C:\\ruta\\a\\la\\carpeta'
convert_all_pdfs_in_folder(folder_path)
